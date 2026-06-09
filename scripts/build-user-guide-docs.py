from __future__ import annotations

import html
import os
import re
from datetime import date
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
EXPORT_DIR = ROOT / "docs" / "user-guides" / "export"
SCREENSHOT_DIR = EXPORT_DIR / "assets" / "screenshots"
PLACEHOLDER_DIR = EXPORT_DIR / "assets" / "placeholders"
GUIDE_IMAGE_DIR = EXPORT_DIR / "assets" / "guide-images"

ACCENT = "006D72"
ACCENT_DARK = "064E52"
YELLOW = "FFD84D"
INK = "0F172A"
MUTED = "64748B"
BORDER = "D7DEE8"
SOFT = "F5F8FB"
CARD = "FFFFFF"
GREEN = "0F766E"
RED = "B42318"
DOC_DATE = os.environ.get("BOXOPS_GUIDE_DATE", date.today().strftime("%d/%m/%Y"))


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_run(paragraph, text: str, *, bold=False, color=INK, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = rgb(color)
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    return paragraph


def add_callout_docx(doc: Document, title: str, text: str, fill: str = "FFF7D6") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [6.25])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "F2D56B")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, text, color=INK)
    doc.add_paragraph()


def add_bullets_docx(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_steps_docx(doc: Document, steps: Iterable[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(step)


def add_route_table_docx(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [1.35, 1.45, 3.45])
    headers = ["Pantalla", "Ruta", "Sirve para"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, ACCENT_DARK)
        set_cell_border(cell, ACCENT_DARK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color="FFFFFF")

    for screen, route, purpose in rows:
        cells = table.add_row().cells
        values = [screen, route, purpose]
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, bold=index == 1, color=INK if index != 1 else ACCENT_DARK)
    doc.add_paragraph()


def add_feature_table_docx(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [1.45, 1.45, 3.35])
    headers = ["Función", "Dónde está", "Para qué sirve"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "F5FBFB")
        set_cell_border(cell, "B9D9DB")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color=ACCENT_DARK)

    for feature, route, purpose in rows:
        cells = table.add_row().cells
        values = [feature, route, purpose]
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, bold=index == 0 or index == 1, color=INK if index != 1 else ACCENT_DARK)
    doc.add_paragraph()


def add_image_docx(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        paragraph.add_run().add_picture(str(image_path), width=Inches(6.25))
    except Exception:
        return
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    run = caption_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(MUTED)


def create_placeholder(name: str, title: str, subtitle: str, route: str) -> Path:
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    path = PLACEHOLDER_DIR / name
    width, height = 1400, 850
    image = Image.new("RGB", (width, height), "#F5F8FB")
    draw = ImageDraw.Draw(image)

    try:
        font_big = ImageFont.truetype("arial.ttf", 56)
        font_medium = ImageFont.truetype("arial.ttf", 30)
        font_small = ImageFont.truetype("arial.ttf", 24)
        font_tiny = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font_big = font_medium = font_small = font_tiny = ImageFont.load_default()

    draw.rounded_rectangle((44, 44, width - 44, height - 44), radius=28, fill="#FFFFFF", outline="#D7DEE8", width=2)
    draw.rounded_rectangle((82, 86, 190, 194), radius=26, fill=f"#{ACCENT}")
    draw.text((118, 116), "B", fill="white", font=font_big, anchor="mm")
    draw.text((220, 92), "BoxOps", fill=f"#{INK}", font=font_medium)
    draw.text((220, 132), "Guía de usuario", fill=f"#{MUTED}", font=font_small)
    draw.rounded_rectangle((82, 244, width - 82, 330), radius=18, fill="#FFF7D6", outline="#F2D56B")
    draw.text((118, 265), route, fill=f"#{ACCENT_DARK}", font=font_small)
    draw.text((118, 382), title, fill=f"#{INK}", font=font_big)
    draw.text((118, 456), subtitle, fill=f"#{475569}", font=font_medium)
    draw.line((118, 552, width - 118, 552), fill="#D7DEE8", width=3)

    card_y = 604
    card_w = 360
    labels = ["Vista principal", "Acciones", "Estados"]
    for index, label in enumerate(labels):
        x = 118 + index * (card_w + 30)
        draw.rounded_rectangle((x, card_y, x + card_w, card_y + 130), radius=20, fill="#F8FAFC", outline="#D7DEE8")
        draw.ellipse((x + 26, card_y + 34, x + 86, card_y + 94), fill="#FFF2B8")
        draw.text((x + 112, card_y + 36), label, fill=f"#{INK}", font=font_small)
        draw.text((x + 112, card_y + 72), "Resumen visual de la función.", fill=f"#{MUTED}", font=font_tiny)

    image.save(path, quality=95)
    return path


def image_ref(name: str, title: str, subtitle: str, route: str) -> Path:
    screenshot = SCREENSHOT_DIR / name
    if screenshot.exists():
        return prepare_guide_image(screenshot)
    return create_placeholder(name, title, subtitle, route)


def prepare_guide_image(source: Path) -> Path:
    GUIDE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    target = GUIDE_IMAGE_DIR / source.name
    source_stat = source.stat()
    if target.exists() and target.stat().st_mtime >= source_stat.st_mtime:
        return target

    image = Image.open(source).convert("RGB")
    width, height = image.size
    max_height = min(height, int(width * 0.72))
    if height > max_height:
        image = image.crop((0, 0, width, max_height))

    image.save(target, quality=96)
    return target


GUIDES = {
    "owner": {
        "role": "Propietario",
        "subtitle": "Control general del box, accesos, configuración y supervisión operativa.",
        "audience": "Para quien toma decisiones sobre el box y necesita ver si la semana está bajo control.",
        "can": [
            "Entrar en Inicio para ver estado semanal, cobertura, avisos y acciones rápidas.",
            "Crear y revisar centros, tipos de actividad, equipo, plantillas y horario.",
            "Gestionar altas de usuario, invitaciones, roles y estado de acceso.",
            "Supervisar solicitudes, ausencias, fichaje, documentos y plan/facturación.",
            "Cambiar ajustes de organización que afectan a la forma de trabajar.",
        ],
        "routes": [
            ("Inicio", "/app", "Estado semanal, avisos y accesos rápidos."),
            ("Equipo", "/app/coaches", "Usuarios, roles, invitaciones y fichas operativas."),
            ("Centros", "/app/centers", "Sedes disponibles para planificar."),
            ("Tipos", "/app/class-types", "Catálogo de actividades y requisitos."),
            ("Plantillas", "/app/templates", "Semana base reutilizable."),
            ("Horario", "/app/schedule", "Bloques reales de la semana y asignaciones."),
            ("Cobertura", "/app/coverage", "Riesgos de clases sin cubrir o insuficientes."),
            ("Configuración", "/app/settings", "Nombre, color y reglas operativas."),
        ],
        "workflows": [
            ("Preparar la base del box", [
                "Crea o revisa los centros en Centros.",
                "Define tipos de actividad: WOD, Open Box, recepción, competición o lo que use el box.",
                "En Equipo, crea cuentas o envía invitaciones y confirma que cada persona tiene rol correcto.",
                "Completa ficha operativa: centro principal, horas semanales y estado operativo.",
            ]),
            ("Crear una semana normal", [
                "En Plantillas, crea una plantilla semanal con los bloques habituales.",
                "Añade bloques por día y hora, con entrenadores necesarios y entrenador por defecto si lo sabes.",
                "Aplica la plantilla a la semana que quieres preparar.",
                "En Horario, revisa excepciones: clases canceladas, cambios de hora o entrenadores distintos.",
            ]),
            ("Cerrar cobertura", [
                "Abre Inicio o Cobertura y localiza riesgos: sin cubrir, insuficientes o conflictos.",
                "Entra al bloque real desde el enlace de la cola.",
                "Añade o retira entrenadores en Asignaciones actuales.",
                "Vuelve a Cobertura y confirma que el riesgo desaparece.",
            ]),
            ("Cuando revises BoxOps", [
                "Anota qué querías hacer, en qué pantalla estabas y qué resultado esperabas.",
                "Si algo no se entiende, no lo des por sabido: ese punto es feedback valioso.",
                "No uses datos laborales sensibles reales todavía.",
            ]),
        ],
        "other_features": [
            ("Cobertura", "/app/coverage", "Reúne clases sin cubrir, cobertura insuficiente y conflictos de asignación para resolverlos desde el bloque real."),
            ("Solicitudes", "/app/requests", "Centraliza peticiones de cambio o cobertura para revisar qué persona pide ayuda y qué bloque se ve afectado."),
            ("Ausencias", "/app/absences", "Permite revisar peticiones y estados de ausencia, y entender su impacto operativo sin mostrar motivos sensibles innecesarios."),
            ("Jornadas previstas", "/app/work-windows", "Define cuándo se espera que una persona esté disponible o presente, separado del horario de clases."),
            ("Mi fichaje", "/app/time", "Registra jornada propia y permite revisar cierres, correcciones y posibles excesos como apoyo interno."),
            ("Mi cuenta y firma", "/app/account", "Guarda perfil propio, avatar privado y firma interna propia. La firma no permite firmar por otra persona."),
            ("Documentos", "/app/documents", "Muestra documentos autorizados y archivos vinculados a programación o contexto interno permitido."),
            ("Estadísticas", "/app/stats", "Resume carga por coach, clases por tipo, actividad por centro y avisos operativos."),
            ("Plan y facturación", "/app/settings/billing", "Muestra plan, límites y uso disponible para la organización."),
        ],
        "images": [
            ("coach-dashboard.png", "Inicio muestra la salud de la semana y accesos rápidos."),
            ("coach-team.png", "Equipo concentra accesos, roles y datos operativos."),
            ("coach-templates.png", "Plantillas permite construir una semana base reutilizable."),
            ("placeholder-coverage.png", "Cobertura: vista para revisar riesgos, huecos y conflictos."),
        ],
        "limits": [
            "No hay reservas de clientes ni BoxWOD todavía.",
            "No hay geolocalización real, payroll ni nómina.",
            "No se deben cargar documentos sensibles reales sin entorno validado.",
            "El acceso de clientes finales y reservas no forma parte de BoxOps.",
        ],
    },
    "admin": {
        "role": "Administrador",
        "subtitle": "Gestión práctica de centros, equipo, catálogo, plantillas, horario y cobertura.",
        "audience": "Para quien configura la operativa diaria y mantiene los datos listos para trabajar.",
        "can": [
            "Crear centros y tipos de actividad.",
            "Crear o invitar usuarios y mantener fichas operativas.",
            "Crear plantillas semanales y aplicarlas a semanas concretas.",
            "Editar bloques del horario y asignar entrenadores.",
            "Revisar cobertura, solicitudes, ausencias y fichajes según permisos activos.",
        ],
        "routes": [
            ("Equipo", "/app/coaches", "Alta de usuarios, invitaciones y fichas."),
            ("Centros", "/app/centers", "Crear o desactivar sedes."),
            ("Tipos", "/app/class-types", "Configurar actividades y requisitos."),
            ("Plantillas", "/app/templates", "Crear patrones semanales."),
            ("Horario", "/app/schedule", "Editar la semana real."),
            ("Cobertura", "/app/coverage", "Resolver huecos de entrenadores."),
            ("Jornadas", "/app/work-windows", "Presencia prevista del personal."),
            ("Mi fichaje", "/app/time", "Fichaje propio y revisión operativa si aplica."),
        ],
        "workflows": [
            ("Dar de alta a una persona", [
                "Entra en Equipo.",
                "Elige Invitar usuario si quieres que acepte por email.",
                "Elige Crear cuenta si vas a darle una contraseña temporal.",
                "Asigna rol, estado y ficha operativa.",
                "Comprueba que queda como cuenta vinculada y operativa.",
            ]),
            ("Crear un centro", [
                "Entra en Centros.",
                "Crea el centro con nombre, slug y zona horaria.",
                "Mantén el centro activo si se va a usar en horario.",
                "Desactívalo, no lo borres, si deja de usarse.",
            ]),
            ("Crear el catálogo de clases", [
                "Entra en Tipos de actividad.",
                "Crea cada actividad con nombre, categoría, color y entrenadores necesarios.",
                "Usa 0 entrenadores solo para bloques que no necesitan cobertura.",
                "Desactiva tipos antiguos en vez de eliminarlos.",
            ]),
            ("Editar el horario", [
                "Abre Horario y elige semana y centro.",
                "Crea bloque manual con el botón + si es algo puntual.",
                "Abre un bloque para cambiar hora, tipo, notas o asignaciones.",
                "Cancela un bloque si no se hará; no lo borres de la historia operativa.",
            ]),
        ],
        "other_features": [
            ("Cobertura", "/app/coverage", "Ayuda a detectar huecos y asignar entrenadores sin ir bloque por bloque a ciegas."),
            ("Solicitudes", "/app/requests", "Sirve para revisar cambios o peticiones de cobertura cuando alguien no puede cubrir un bloque."),
            ("Ausencias", "/app/absences", "Permite ver ausencias y estados que pueden afectar a horario y cobertura."),
            ("Jornadas previstas", "/app/work-windows", "Gestiona presencia prevista del equipo sin convertirla en clase ni asignación."),
            ("Mi fichaje", "/app/time", "Permite fichar, revisar la semana y gestionar correcciones según la política activa."),
            ("Mi cuenta y firma", "/app/account", "Permite actualizar datos propios, avatar y firma interna privada."),
            ("Documentos", "/app/documents", "Muestra documentos permitidos y adjuntos internos autorizados."),
            ("Estadísticas", "/app/stats", "Ofrece lectura operativa de carga, centros, tipos y avisos."),
        ],
        "images": [
            ("coach-team.png", "Equipo: altas, roles y fichas operativas."),
            ("coach-centers.png", "Centros: sedes y estado operativo."),
            ("coach-class-types.png", "Tipos de actividad: catálogo que alimenta horarios y plantillas."),
            ("coach-templates.png", "Plantillas: bloques base y entrenadores por defecto."),
        ],
        "limits": [
            "No uses el rol de administrador para acceder a datos sensibles que no estén diseñados para ese permiso.",
            "No prometas automatismos: cobertura, ausencias y eventos siguen requiriendo revisión humana.",
            "No cargues datos reales sensibles hasta cerrar entorno, permisos y validación.",
        ],
    },
    "manager": {
        "role": "Responsable",
        "subtitle": "Supervisión diaria: horario, cobertura, solicitudes, ausencias, jornadas y fichaje.",
        "audience": "Para quien resuelve el día a día sin tocar la configuración global del negocio.",
        "can": [
            "Ver Inicio y detectar qué necesita atención.",
            "Gestionar bloques, asignaciones y cobertura operativa.",
            "Revisar solicitudes de cambio o cobertura cuando existan.",
            "Revisar ausencias y su impacto operativo.",
            "Consultar fichaje y posibles excesos como apoyo interno, sin payroll.",
        ],
        "routes": [
            ("Inicio", "/app", "Vista rápida de semana y pendientes."),
            ("Horario", "/app/schedule", "Bloques, asignaciones y contexto."),
            ("Cobertura", "/app/coverage", "Lista de riesgos accionables."),
            ("Solicitudes", "/app/requests", "Cambios o cobertura pendiente."),
            ("Ausencias", "/app/absences", "Peticiones y estados de ausencia."),
            ("Jornadas", "/app/work-windows", "Presencia prevista."),
            ("Mi fichaje", "/app/time", "Registro propio y revisión si aplica."),
            ("Estadísticas", "/app/stats", "Lectura operativa de carga y avisos."),
        ],
        "workflows": [
            ("Revisar la semana", [
                "Empieza en Inicio.",
                "Mira cobertura semanal, avisos y accesos rápidos.",
                "Abre Horario para entender el contexto por día y centro.",
                "Abre Cobertura si hay riesgos que resolver.",
            ]),
            ("Resolver una clase sin cubrir", [
                "En Cobertura, abre el riesgo concreto.",
                "Comprueba centro, hora, tipo de actividad y entrenadores necesarios.",
                "Asigna un entrenador disponible o ajusta el requisito si el bloque está mal configurado.",
                "Revisa que no aparezcan conflictos de solape.",
            ]),
            ("Usar solicitudes y ausencias", [
                "En Solicitudes, revisa qué pide la persona y qué bloque afecta.",
                "En Ausencias, revisa fechas y estado sin exponer motivos sensibles innecesarios.",
                "Vuelve a Horario o Cobertura para comprobar impacto.",
            ]),
            ("Fichaje y excesos", [
                "En Mi fichaje, revisa registros propios y, si tu rol lo permite, correcciones del equipo.",
                "Usa posibles excesos como alerta operativa, no como aprobación legal ni nómina.",
                "Documenta cualquier caso raro para revisarlo antes de beta real.",
            ]),
        ],
        "other_features": [
            ("Cobertura", "/app/coverage", "Prioriza clases sin cubrir, insuficientes o con conflicto para resolver la semana con menos vueltas."),
            ("Solicitudes", "/app/requests", "Muestra peticiones de cambio o cobertura y ayuda a decidir qué hacer con cada caso."),
            ("Ausencias", "/app/absences", "Ayuda a revisar ausencias y comprobar si afectan a bloques o cobertura."),
            ("Jornadas previstas", "/app/work-windows", "Da contexto de disponibilidad/presencia prevista para organizar mejor el día."),
            ("Mi fichaje", "/app/time", "Registra jornada propia y, si el rol lo permite, revisa correcciones o cierres del equipo."),
            ("Mi cuenta y firma", "/app/account", "Gestiona perfil propio, avatar y firma interna privada."),
            ("Documentos", "/app/documents", "Permite consultar documentos o programación autorizada sin abrir información sensible por defecto."),
            ("Estadísticas", "/app/stats", "Resume actividad, carga y avisos para entender la operación sin entrar a cada pantalla."),
        ],
        "images": [
            ("coach-dashboard.png", "Inicio ayuda a priorizar la semana."),
            ("placeholder-coverage.png", "Cobertura: vista para priorizar riesgos de la semana."),
            ("coach-time.png", "Mi fichaje muestra registro, balance semanal y revisión operativa."),
            ("placeholder-requests-absences.png", "Solicitudes y ausencias: seguimiento de peticiones y estados."),
        ],
        "limits": [
            "No cambia planes, facturación ni configuración global.",
            "No debe gestionar roles o cuentas si la acción no aparece en Equipo.",
            "No hay payroll ni aprobación legal definitiva de horas extra.",
            "No hay resolución automática de cobertura: la decisión sigue siendo humana.",
        ],
    },
    "coach": {
        "role": "Entrenador",
        "subtitle": "Consulta de horario, equipo, cuenta propia, fichaje, solicitudes y documentos permitidos.",
        "audience": "Para quien necesita saber cuándo trabaja, registrar jornada y pedir cambios sin tocar la gestión del box.",
        "can": [
            "Entrar y consultar su contexto de trabajo.",
            "Ver horario semanal y usar Mi horario si tiene ficha vinculada.",
            "Consultar equipo, centros, tipos de actividad y plantillas en lectura.",
            "Registrar entrada/salida manual y revisar la semana de fichaje.",
            "Gestionar su perfil, avatar y firma interna propia.",
        ],
        "routes": [
            ("Inicio", "/app", "Avisos propios y accesos rápidos."),
            ("Horario", "/app/schedule", "Semana visible y filtro Mi horario."),
            ("Equipo", "/app/coaches", "Consulta del equipo visible."),
            ("Mi fichaje", "/app/time", "Entrada, salida, semana y correcciones."),
            ("Ausencias", "/app/absences", "Peticiones propias si están activas."),
            ("Solicitudes", "/app/requests", "Cambios o cobertura cuando aplique."),
            ("Documentos", "/app/documents", "Documentos permitidos por acceso."),
            ("Mi cuenta", "/app/account", "Perfil, avatar y firma propia."),
        ],
        "workflows": [
            ("Entrar y orientarte", [
                "Abre la URL de BoxOps y entra con tu email y contraseña.",
                "Si te piden cambiar contraseña, hazlo antes de seguir.",
                "Empieza en Inicio para ver avisos y accesos directos.",
                "En móvil, usa la navegación inferior: Inicio, Horario, Equipo y Más.",
            ]),
            ("Consultar tu horario", [
                "Entra en Horario.",
                "Elige la semana con Anterior, Hoy o Siguiente.",
                "Activa Mi horario si aparece disponible.",
                "Abre un bloque para ver centro, hora, actividad, notas y asignaciones.",
            ]),
            ("Fichar manualmente", [
                "Entra en Mi fichaje.",
                "Revisa fecha, hora y centro si aplica.",
                "Pulsa Fichar entrada al empezar y Fichar salida al terminar.",
                "Si te equivocas, usa correcciones con motivo claro.",
            ]),
            ("Mantener tu cuenta", [
                "En Mi cuenta, revisa tu perfil visible.",
                "Actualiza nombre visible, alias o email público si está permitido.",
                "Sube tu avatar propio si se usa en el piloto.",
                "Guarda tu firma interna solo como confirmación propia; aún no firma documentos.",
            ]),
        ],
        "other_features": [
            ("Cobertura", "/app/schedule", "En Horario puedes ver si un bloque está cubierto, sin cubrir, insuficiente o en conflicto."),
            ("Solicitudes", "/app/requests", "Sirve para pedir o revisar cambios/coberturas propias cuando el flujo esté disponible para tu rol."),
            ("Ausencias", "/app/absences", "Permite consultar o crear peticiones propias de ausencia si la organización lo usa."),
            ("Mi fichaje", "/app/time", "Registra entrada/salida manual, revisa la semana y pide correcciones si algo no cuadra."),
            ("Mi cuenta y firma", "/app/account", "Permite actualizar tu perfil visible, avatar y firma interna propia."),
            ("Documentos", "/app/documents", "Muestra documentos disponibles para ti según permisos."),
            ("Equipo, centros y tipos", "/app/more", "Permiten consultar información operativa básica en modo lectura."),
            ("Plantillas", "/app/templates", "Ayudan a entender la estructura semanal base, aunque no puedas editarlas."),
        ],
        "images": [
            ("login.png", "Login: punto de entrada a BoxOps."),
            ("coach-mobile-home.png", "Vista móvil: navegación pensada para uso rápido."),
            ("coach-schedule.png", "Horario: semana, filtros y contexto de bloques."),
            ("coach-time.png", "Mi fichaje: entrada/salida y balance semanal."),
            ("account.png", "Mi cuenta: perfil propio, avatar y firma interna."),
        ],
        "limits": [
            "No crea ni edita centros, tipos, usuarios, plantillas ni bloques.",
            "No ve datos laborales sensibles por defecto.",
            "No hay geolocalización real ni tracking de ubicación.",
            "No hay reservas de clientes en BoxOps; eso será BoxWOD si se construye.",
        ],
    },
}


def html_list(items: Iterable[str]) -> str:
    return "".join(f"<li>{html.escape(item)}</li>" for item in items)


def html_steps(items: Iterable[str]) -> str:
    return "".join(f"<li>{html.escape(item)}</li>" for item in items)


def image_src(path: Path) -> str:
    return path.relative_to(EXPORT_DIR).as_posix()


def build_html(key: str, guide: dict) -> str:
    title = f"BoxOps - Guía de {guide['role']}"
    images = []
    for name, caption in guide["images"]:
        path = image_ref(
            name,
            caption,
            "Referencia visual de la pantalla.",
            "/app",
        )
        images.append((path, caption))

    route_rows = "\n".join(
        "<tr>"
        f"<td>{html.escape(screen)}</td>"
        f"<td><code>{html.escape(route)}</code></td>"
        f"<td>{html.escape(purpose)}</td>"
        "</tr>"
        for screen, route, purpose in guide["routes"]
    )

    other_feature_rows = "\n".join(
        "<tr>"
        f"<td>{html.escape(feature)}</td>"
        f"<td><code>{html.escape(route)}</code></td>"
        f"<td>{html.escape(purpose)}</td>"
        "</tr>"
        for feature, route, purpose in guide["other_features"]
    )

    workflow_html = []
    for workflow_title, steps in guide["workflows"]:
        workflow_html.append(
            f"""
            <section class="workflow">
              <h3>{html.escape(workflow_title)}</h3>
              <ol>{html_steps(steps)}</ol>
            </section>
            """
        )

    image_html = "\n".join(
        f"""
        <figure>
          <img src="{image_src(path)}" alt="{html.escape(caption)}">
          <figcaption>{html.escape(caption)}</figcaption>
        </figure>
        """
        for path, caption in images
    )

    return f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>
    @page {{ size: Letter; margin: 0.72in; }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: Arial, Helvetica, sans-serif;
      color: #{INK};
      background: white;
      font-size: 10.5pt;
      line-height: 1.45;
    }}
    .cover {{
      min-height: 9.2in;
      display: flex;
      flex-direction: column;
      justify-content: space-between;
      border: 1px solid #{BORDER};
      border-radius: 18px;
      padding: 42px;
      background:
        linear-gradient(145deg, rgba(0,109,114,.10), rgba(255,216,77,.16) 45%, rgba(255,255,255,1) 76%),
        #fff;
    }}
    .brand {{ display: flex; align-items: center; gap: 14px; }}
    .mark {{
      width: 46px; height: 46px; border-radius: 14px;
      background: #{ACCENT}; color: white;
      display: grid; place-items: center;
      font-weight: 800; font-size: 20px;
    }}
    .eyebrow {{
      color: #{ACCENT_DARK}; font-weight: 700; letter-spacing: .08em;
      text-transform: uppercase; font-size: 10px;
    }}
    h1 {{ font-size: 38px; line-height: 1.06; margin: 72px 0 12px; letter-spacing: 0; }}
    .subtitle {{ font-size: 17px; color: #334155; max-width: 620px; }}
    .meta-strip {{
      display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px;
      margin-top: 36px;
    }}
    .meta-card {{
      border: 1px solid #{BORDER}; border-radius: 14px; padding: 14px 16px;
      background: rgba(255,255,255,.76);
    }}
    .meta-card b {{ display: block; font-size: 9px; color: #{MUTED}; text-transform: uppercase; letter-spacing: .06em; }}
    .meta-card span {{ display: block; margin-top: 4px; font-weight: 700; }}
    .notice {{
      margin-top: 26px; border: 1px solid #F2D56B; border-radius: 14px;
      background: #FFF7D6; padding: 15px 18px; color: #3B3A22;
    }}
    .page {{ break-before: page; }}
    h2 {{ font-size: 23px; margin: 0 0 12px; color: #{ACCENT_DARK}; }}
    h3 {{ font-size: 15px; margin: 18px 0 8px; color: #{INK}; }}
    p {{ margin: 0 0 9px; }}
    .lead {{ color: #475569; font-size: 12pt; margin-bottom: 18px; }}
    .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }}
    .card {{
      border: 1px solid #{BORDER}; border-radius: 14px; padding: 15px 16px;
      background: #{CARD}; break-inside: avoid;
    }}
    .card h3 {{ margin-top: 0; }}
    ul, ol {{ margin: 8px 0 0 21px; padding: 0; }}
    li {{ margin: 0 0 5px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 14px 0 20px; break-inside: avoid; }}
    th {{ text-align: left; background: #{ACCENT_DARK}; color: white; padding: 9px 10px; font-size: 9.6pt; }}
    td {{ border: 1px solid #{BORDER}; padding: 9px 10px; vertical-align: top; }}
    code {{ color: #{ACCENT_DARK}; font-weight: 700; }}
    .workflow {{
      border: 1px solid #{BORDER}; border-radius: 14px; padding: 14px 18px;
      background: #FAFCFE; margin: 0 0 12px; break-inside: avoid;
    }}
    .workflow h3 {{ margin-top: 0; }}
    figure {{
      margin: 18px 0 24px; break-inside: avoid;
      border: 1px solid #{BORDER}; border-radius: 16px;
      overflow: hidden; background: white;
    }}
    figure img {{ display: block; width: 100%; max-height: 6.8in; object-fit: contain; background: #F5F8FB; }}
    figcaption {{
      padding: 9px 12px; color: #{MUTED}; font-size: 9.3pt;
      border-top: 1px solid #{BORDER}; background: #F8FAFC;
    }}
    .limits {{
      border-left: 5px solid #{ACCENT}; background: #F5FBFB;
      padding: 14px 18px; border-radius: 12px; margin-top: 10px;
    }}
    .footer-note {{ color: #{MUTED}; font-size: 9pt; margin-top: 18px; }}
  </style>
</head>
<body>
  <section class="cover">
    <div>
      <div class="brand">
        <div class="mark">B</div>
        <div>
          <div class="eyebrow">BoxOps</div>
          <div>Guía de usuario</div>
        </div>
      </div>
      <h1>Guía de {html.escape(guide['role'])}</h1>
      <p class="subtitle">{html.escape(guide['subtitle'])}</p>
      <div class="meta-strip">
        <div class="meta-card"><b>Versión</b><span>v1</span></div>
        <div class="meta-card"><b>Fecha</b><span>{DOC_DATE}</span></div>
        <div class="meta-card"><b>Rol</b><span>{html.escape(guide['role'])}</span></div>
      </div>
    </div>
    <p class="footer-note">Guía rápida para trabajar con BoxOps sin perder contexto.</p>
  </section>

  <section class="page">
    <h2>Para quién es esta guía</h2>
    <p class="lead">{html.escape(guide['audience'])}</p>
    <div class="grid">
      <div class="card">
        <h3>Qué puedes hacer</h3>
        <ul>{html_list(guide['can'])}</ul>
      </div>
      <div class="card">
        <h3>Uso recomendado</h3>
        <ul>
          <li>Empieza por las pantallas principales de tu rol.</li>
          <li>Si una acción no aparece, puede depender de tus permisos.</li>
          <li>No des por hecho que un botón “debería estar”: si no aparece, puede ser por rol.</li>
          <li>Reporta página, acción, resultado esperado y resultado real.</li>
        </ul>
      </div>
    </div>

    <h2 style="margin-top:24px">Mapa rápido</h2>
    <table>
      <thead><tr><th>Pantalla</th><th>Ruta</th><th>Sirve para</th></tr></thead>
      <tbody>{route_rows}</tbody>
    </table>
  </section>

  <section class="page">
    <h2>Flujos principales</h2>
    <p class="lead">Empieza por estos recorridos. Son los que más valor dan en el uso diario.</p>
    {''.join(workflow_html)}
  </section>

  <section class="page">
    <h2>Otras funciones útiles</h2>
    <p class="lead">Estas zonas no siempre son la primera parada, pero ayudan a entender la operación completa.</p>
    <table>
      <thead><tr><th>Función</th><th>Dónde está</th><th>Para qué sirve</th></tr></thead>
      <tbody>{other_feature_rows}</tbody>
    </table>
  </section>

  <section class="page">
    <h2>Capturas de referencia</h2>
    <p class="lead">Las imágenes ayudan a reconocer la pantalla y los elementos principales.</p>
    {image_html}
  </section>

  <section class="page">
    <h2>Límites actuales</h2>
    <div class="limits">
      <ul>{html_list(guide['limits'])}</ul>
    </div>
    <h2 style="margin-top:26px">Si necesitas reportar algo</h2>
    <ol>
      <li>Di qué estabas intentando hacer.</li>
      <li>Indica la pantalla: por ejemplo, Horario, Equipo o Mi fichaje.</li>
      <li>Explica qué esperabas que pasara.</li>
      <li>Explica qué pasó realmente.</li>
      <li>Añade captura si el problema es visual o de comprensión.</li>
    </ol>
    <p class="footer-note">Evita enviar contraseñas, datos sensibles o información laboral real por canales no acordados.</p>
  </section>
</body>
</html>"""


def setup_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 18, ACCENT_DARK, 18, 8),
        ("Heading 2", 14, ACCENT_DARK, 13, 6),
        ("Heading 3", 12, INK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(10.5)
    styles["List Number"].font.name = "Calibri"
    styles["List Number"].font.size = Pt(10.5)


def add_docx_cover(doc: Document, guide: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "BoxOps", bold=True, color=ACCENT_DARK, size=18)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(28)
    add_run(p2, "Guía de usuario", color=MUTED, size=11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    add_run(title, f"Guía de {guide['role']}", bold=True, color=INK, size=30)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    add_run(subtitle, guide["subtitle"], color="334155", size=13)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [2.05, 2.05, 2.05])
    values = [("Versión", "v1"), ("Fecha", DOC_DATE), ("Rol", guide["role"])]
    for index, (label, value) in enumerate(values):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "FFFFFF")
        set_cell_border(cell)
        set_cell_margins(cell, top=130, bottom=130)
        add_run(cell.paragraphs[0], label.upper(), bold=True, color=MUTED, size=8)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_run(p, value, bold=True, color=INK, size=11)

    doc.add_paragraph()
    doc.add_page_break()


def build_docx(key: str, guide: dict) -> Path:
    doc = Document()
    setup_docx_styles(doc)
    add_docx_cover(doc, guide)

    doc.add_heading("Para quién es esta guía", level=1)
    p = doc.add_paragraph()
    add_run(p, guide["audience"], color="475569", size=12)

    doc.add_heading("Qué puedes hacer", level=2)
    add_bullets_docx(doc, guide["can"])

    doc.add_heading("Mapa rápido", level=2)
    add_route_table_docx(doc, guide["routes"])

    doc.add_heading("Flujos principales", level=1)
    for title, steps in guide["workflows"]:
        doc.add_heading(title, level=2)
        add_steps_docx(doc, steps)

    doc.add_heading("Otras funciones útiles", level=1)
    p = doc.add_paragraph()
    add_run(
        p,
        "Estas zonas no siempre son la primera parada, pero ayudan a entender la operación completa.",
        color=MUTED,
    )
    add_feature_table_docx(doc, guide["other_features"])

    doc.add_page_break()
    doc.add_heading("Capturas de referencia", level=1)
    note = doc.add_paragraph()
    add_run(
        note,
        "Las capturas ayudan a reconocer pantallas y elementos principales.",
        color=MUTED,
    )
    for name, caption in guide["images"]:
        image_path = image_ref(
            name,
            caption,
            "Referencia visual de la pantalla.",
            "/app",
        )
        add_image_docx(doc, image_path, caption)

    doc.add_heading("Límites actuales", level=1)
    add_bullets_docx(doc, guide["limits"])

    doc.add_heading("Si necesitas reportar algo", level=1)
    add_steps_docx(
        doc,
        [
            "Di qué estabas intentando hacer.",
            "Indica la pantalla: Horario, Equipo, Mi fichaje u otra.",
            "Explica qué esperabas que pasara.",
            "Explica qué pasó realmente.",
            "Añade captura si el problema es visual o de comprensión.",
        ],
    )
    add_callout_docx(
        doc,
        "Privacidad",
        "Evita enviar contraseñas, datos sensibles o información laboral real por canales no acordados.",
        fill="F5FBFB",
    )

    path = EXPORT_DIR / f"boxops-guia-{slugify(guide['role'])}.docx"
    doc.save(path)
    return path


def main() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    GUIDE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    generated = []
    for key, guide in GUIDES.items():
        html_path = EXPORT_DIR / f"boxops-guia-{slugify(guide['role'])}.html"
        html_path.write_text(build_html(key, guide), encoding="utf-8")
        docx_path = build_docx(key, guide)
        generated.append(str(html_path.relative_to(ROOT)))
        generated.append(str(docx_path.relative_to(ROOT)))

    report = "\n".join(generated)
    (EXPORT_DIR / "build-report.txt").write_text(report + "\n", encoding="utf-8")
    print(report)


if __name__ == "__main__":
    main()
